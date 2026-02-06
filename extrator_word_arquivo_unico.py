"""
Extrator de Fichas Word - Modo Arquivo Único (Página por Página)
Versão 1.1.0

Processa um único arquivo Word com múltiplas fichas separadas por quebras de página.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import threading
from datetime import datetime
import docx
import pandas as pd
from typing import Dict, List

# Importar a classe base do extrator original
import sys
sys.path.append(os.path.dirname(__file__))

class ExtratorWordArquivoUnico:
    """Interface gráfica para extração de arquivo único com múltiplas páginas"""
    
    VERSION = "1.1.0"
    
    def __init__(self, root):
        self.root = root
        self.root.title(f"Extrator de Fichas Word v{self.VERSION} - Arquivo Único")
        self.root.geometry("750x550")
        self.root.resizable(True, True)
        
        # Variáveis
        self.arquivo_selecionado = tk.StringVar()
        self.total_paginas = 0
        self.paginas_processadas = 0
        self.processando = False
        
        # Configurar estilo
        self.configurar_estilo()
        
        # Criar interface
        self.criar_interface()
        
        # Centralizar janela
        self.centralizar_janela()
    
    def configurar_estilo(self):
        """Configura o estilo visual moderno com tema escuro"""
        style = ttk.Style()
        style.theme_use('clam')
        
        # Cores do tema escuro moderno
        self.cor_fundo = "#1e1e2e"
        self.cor_fundo_sec = "#2a2a3e"
        self.cor_primaria = "#89b4fa"
        self.cor_secundaria = "#74c7ec"
        self.cor_sucesso = "#a6e3a1"
        self.cor_aviso = "#f9e2af"
        self.cor_erro = "#f38ba8"
        self.cor_texto = "#cdd6f4"
        self.cor_texto_sec = "#a6adc8"
        self.cor_borda = "#45475a"
        self.cor_botao = "#313244"
        self.cor_botao_hover = "#45475a"
        
        # Configurar estilos
        style.configure('Title.TLabel', font=('Segoe UI', 14, 'bold'),
                       foreground=self.cor_primaria, background=self.cor_fundo)
        style.configure('Subtitle.TLabel', font=('Segoe UI', 9),
                       foreground=self.cor_texto_sec, background=self.cor_fundo)
        style.configure('TLabel', foreground=self.cor_texto, background=self.cor_fundo)
        style.configure('TFrame', background=self.cor_fundo)
        style.configure('TLabelframe', background=self.cor_fundo,
                       foreground=self.cor_texto, bordercolor=self.cor_borda)
        style.configure('TLabelframe.Label', background=self.cor_fundo,
                       foreground=self.cor_primaria, font=('Segoe UI', 10, 'bold'))
        style.configure('TEntry', fieldbackground=self.cor_fundo_sec,
                       foreground=self.cor_texto, bordercolor=self.cor_borda)
        style.configure('TProgressbar', background=self.cor_primaria,
                       troughcolor=self.cor_fundo_sec, bordercolor=self.cor_borda,
                       lightcolor=self.cor_primaria, darkcolor=self.cor_primaria)
        
        self.root.configure(bg=self.cor_fundo)
    
    def criar_interface(self):
        """Cria todos os componentes da interface"""
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Cabeçalho
        self.criar_cabecalho(main_frame)
        
        # Seleção de arquivo
        self.criar_secao_arquivo(main_frame)
        
        # Progresso
        self.criar_secao_progresso(main_frame)
        
        # Log
        self.criar_secao_log(main_frame)
        
        # Botões
        self.criar_secao_botoes(main_frame)
        
        # Rodapé
        self.criar_rodape(main_frame)
    
    def criar_cabecalho(self, parent):
        """Cria o cabeçalho"""
        header_frame = ttk.Frame(parent)
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(header_frame, text="📄 Extrator de Fichas Word",
                 style='Title.TLabel').pack()
        ttk.Label(header_frame, text="Modo: Arquivo Único (Página por Página)",
                 style='Subtitle.TLabel').pack()
        ttk.Label(header_frame, text=f"v{self.VERSION}",
                 font=('Segoe UI', 8), foreground='#64748b',
                 background=self.cor_fundo).pack(pady=(5, 0))
    
    def criar_secao_arquivo(self, parent):
        """Cria a seção de seleção de arquivo"""
        file_frame = ttk.LabelFrame(parent, text="📄 Arquivo Word", padding="8")
        file_frame.pack(fill=tk.X, pady=(0, 8))
        
        self.file_entry = ttk.Entry(file_frame, textvariable=self.arquivo_selecionado,
                                    state='readonly', font=('Segoe UI', 8))
        self.file_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        btn_selecionar = tk.Button(file_frame, text="📄 Selecionar",
                                   command=self.selecionar_arquivo,
                                   bg=self.cor_botao, fg=self.cor_texto,
                                   font=('Segoe UI', 9, 'bold'), relief=tk.FLAT,
                                   padx=12, pady=6, cursor='hand2',
                                   activebackground=self.cor_botao_hover,
                                   activeforeground=self.cor_texto)
        btn_selecionar.pack(side=tk.RIGHT)
        btn_selecionar.bind('<Enter>', lambda e: btn_selecionar.config(bg=self.cor_botao_hover))
        btn_selecionar.bind('<Leave>', lambda e: btn_selecionar.config(bg=self.cor_botao))
    
    def criar_secao_progresso(self, parent):
        """Cria a seção de progresso"""
        progress_frame = ttk.LabelFrame(parent, text="📊 Progresso", padding="8")
        progress_frame.pack(fill=tk.X, pady=(0, 8))
        
        self.status_label = ttk.Label(progress_frame, text="Aguardando seleção de arquivo...",
                                     font=('Segoe UI', 8, 'bold'),
                                     foreground=self.cor_texto, background=self.cor_fundo)
        self.status_label.pack(fill=tk.X, pady=(0, 5))
        
        self.progress_bar = ttk.Progressbar(progress_frame, mode='determinate', length=300)
        self.progress_bar.pack(fill=tk.X, pady=(0, 3))
        
        self.percent_label = ttk.Label(progress_frame, text="0%",
                                      font=('Segoe UI', 8),
                                      foreground=self.cor_texto, background=self.cor_fundo)
        self.percent_label.pack()
    
    def criar_secao_log(self, parent):
        """Cria a seção de log"""
        log_frame = ttk.LabelFrame(parent, text="📋 Log", padding="8")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 8))
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=10,
                                                 font=('Consolas', 8),
                                                 bg=self.cor_fundo_sec, fg=self.cor_texto,
                                                 relief=tk.FLAT, borderwidth=1,
                                                 insertbackground=self.cor_texto)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        self.log_text.tag_config('info', foreground=self.cor_secundaria)
        self.log_text.tag_config('success', foreground=self.cor_sucesso)
        self.log_text.tag_config('warning', foreground=self.cor_aviso)
        self.log_text.tag_config('error', foreground=self.cor_erro)
        self.log_text.tag_config('header', foreground=self.cor_primaria, font=('Consolas', 9, 'bold'))
    
    def criar_secao_botoes(self, parent):
        """Cria a seção de botões"""
        btn_frame = ttk.Frame(parent)
        btn_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.btn_processar = tk.Button(btn_frame, text="▶ Processar",
                                      command=self.iniciar_processamento,
                                      bg=self.cor_primaria, fg="#1e1e2e",
                                      font=('Segoe UI', 9, 'bold'), relief=tk.FLAT,
                                      padx=15, pady=6, cursor='hand2',
                                      activebackground=self.cor_secundaria,
                                      activeforeground="#1e1e2e")
        self.btn_processar.pack(side=tk.LEFT, padx=(0, 10))
        self.btn_processar.bind('<Enter>', lambda e: self.btn_processar.config(bg=self.cor_secundaria) if self.btn_processar['state'] != 'disabled' else None)
        self.btn_processar.bind('<Leave>', lambda e: self.btn_processar.config(bg=self.cor_primaria) if self.btn_processar['state'] != 'disabled' else None)
        
        btn_limpar = tk.Button(btn_frame, text="🗑️ Limpar",
                              command=self.limpar_log,
                              bg=self.cor_botao, fg=self.cor_texto,
                              font=('Segoe UI', 9, 'bold'), relief=tk.FLAT,
                              padx=12, pady=6, cursor='hand2',
                              activebackground=self.cor_botao_hover,
                              activeforeground=self.cor_texto)
        btn_limpar.pack(side=tk.LEFT, padx=(0, 10))
        btn_limpar.bind('<Enter>', lambda e: btn_limpar.config(bg=self.cor_botao_hover))
        btn_limpar.bind('<Leave>', lambda e: btn_limpar.config(bg=self.cor_botao))
        
        btn_sair = tk.Button(btn_frame, text="✖ Sair", command=self.sair,
                            bg=self.cor_erro, fg="#1e1e2e",
                            font=('Segoe UI', 9, 'bold'), relief=tk.FLAT,
                            padx=12, pady=6, cursor='hand2',
                            activebackground="#eba0ac", activeforeground="#1e1e2e")
        btn_sair.pack(side=tk.RIGHT)
        btn_sair.bind('<Enter>', lambda e: btn_sair.config(bg="#eba0ac"))
        btn_sair.bind('<Leave>', lambda e: btn_sair.config(bg=self.cor_erro))
    
    def criar_rodape(self, parent):
        """Cria o rodapé"""
        footer_frame = ttk.Frame(parent)
        footer_frame.pack(fill=tk.X)
        ttk.Label(footer_frame, text="Desenvolvido com ❤️ por Antigravity AI",
                 font=('Segoe UI', 8), foreground='#94a3b8',
                 background=self.cor_fundo).pack()
    
    def centralizar_janela(self):
        """Centraliza a janela na tela"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
    
    def adicionar_log(self, mensagem, tag='info'):
        """Adiciona mensagem ao log"""
        timestamp = datetime.now().strftime('%H:%M:%S')
        self.log_text.insert(tk.END, f"[{timestamp}] ", 'header')
        self.log_text.insert(tk.END, f"{mensagem}\n", tag)
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def limpar_log(self):
        """Limpa o log"""
        self.log_text.delete(1.0, tk.END)
        self.adicionar_log("Log limpo.", 'info')
    
    def selecionar_arquivo(self):
        """Abre diálogo para selecionar arquivo"""
        arquivo = filedialog.askopenfilename(
            title="Selecione o arquivo Word com múltiplas fichas",
            filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")]
        )
        
        if arquivo:
            self.arquivo_selecionado.set(arquivo)
            self.adicionar_log(f"Arquivo selecionado: {os.path.basename(arquivo)}", 'success')
            
            # Analisar arquivo
            try:
                doc = docx.Document(arquivo)
                # Contar quebras de página
                quebras = sum(1 for para in doc.paragraphs 
                            if para._element.xpath('.//w:br[@w:type="page"]'))
                self.total_paginas = quebras + 1 if quebras > 0 else 1
                
                self.adicionar_log(f"Páginas detectadas: {self.total_paginas}", 'info')
                self.btn_processar.config(state='normal')
                self.status_label.config(text=f"Pronto para processar {self.total_paginas} página(s)")
            except Exception as e:
                self.adicionar_log(f"Erro ao analisar arquivo: {str(e)}", 'error')
                self.btn_processar.config(state='disabled')
    
    def iniciar_processamento(self):
        """Inicia o processamento"""
        if not self.arquivo_selecionado.get():
            messagebox.showwarning("Aviso", "Selecione um arquivo primeiro!")
            return
        
        self.btn_processar.config(state='disabled')
        self.processando = True
        
        thread = threading.Thread(target=self.processar_arquivo, daemon=True)
        thread.start()
    
    def extrair_todas_as_fichas(self, caminho_arquivo: str) -> List[Dict[str, str]]:
        """Extrai dados de todas as fichas dentro de um documento Word"""
        campos_mapeamento = {
            'Código': 'codigo', 'Contrato': 'contrato', 'Nome do(a) trabalhador(a)': 'nome',
            'Matricula eSocial': 'matricula_esocial', 'Nome do pai': 'nome_pai',
            'Nome da mãe': 'nome_mae', 'Data de nascimento': 'data_nascimento',
            'Raça/cor': 'raca_cor', 'Sexo': 'sexo', 'Naturalidade': 'naturalidade',
            'Nacionalidade': 'nacionalidade', 'Estado Civil': 'estado_civil',
            'Deficiente': 'deficiente', 'Tipo de deficiência': 'tipo_deficiencia',
            'Tipo sanguíneo': 'tipo_sanguineo', 'CPF': 'cpf',
            'Cédula de identidade': 'rg', 'Data de emissão': 'data_emissao_rg',
            'Órgão/UF': 'orgao_uf_rg', 'CTPS': 'ctps', 'Série': 'serie_ctps',
            'Dígito': 'digito_ctps', 'Nº título de eleitor': 'titulo_eleitor',
            'Zona': 'zona_eleitoral', 'Seção': 'secao_eleitoral', 'Nº do PIS': 'pis',
            'Data de cadastramento': 'data_cadastramento_pis', 'Grau de instrução': 'grau_instrucao',
            'Endereço': 'endereco', 'Número': 'numero', 'Complemento': 'complemento',
            'Bairro': 'bairro', 'Cidade': 'cidade', 'Estado': 'estado', 'CEP': 'cep',
            'Telefone': 'telefone', 'Celular': 'celular', 'Endereço eletrônico': 'email',
            'Data de admissão': 'data_admissao', 'Data do registro': 'data_registro',
            'Função': 'funcao', 'CBO': 'cbo', 'Salário Inicial': 'salario_inicial',
            'Forma de pagamento': 'forma_pagamento', 'Tipo de pagamento': 'tipo_pagamento',
            'Insalubridade': 'insalubridade', 'Periculosidade': 'periculosidade',
            'Sindicato': 'sindicato', 'Centro de custo': 'centro_custo',
            'Localização': 'localizacao', 'Horário': 'horario',
            'Nº da conta FGTS': 'conta_fgts', 'Data de opção': 'data_opcao_fgts',
            'Banco depositário - FGTS': 'banco_fgts', 'Data rescisão': 'data_rescisao',
            'Aviso prévio': 'aviso_previo', 'Saldo FGTS': 'saldo_fgts',
            'Maior remuneração': 'maior_remuneracao', 'Causa da rescisão': 'causa_rescisao',
            'Empregador': 'empregador', 'CNPJ': 'cnpj_empregador'
        }
        
        campos_endereco = ['endereco', 'numero', 'complemento', 'bairro', 'cidade', 'estado', 'cep', 'telefone', 'celular']
        
        doc = docx.Document(caminho_arquivo)
        todas_fichas = []
        dados_atuais = {}
        contador_fichas = 0
        
        # Processar todas as tabelas
        for tabela in doc.tables:
            # Ao encontrar uma nova tabela, se ela tiver 'Código' ou 'Nome' na primeira linha, 
            # e já tivermos dados, pode ser o início de uma nova ficha
            for row_idx, row in enumerate(tabela.rows):
                for cell in row.cells:
                    texto_celula = cell.text.strip()
                    
                    if not texto_celula: continue

                    # Verificar se é o início de uma nova ficha (Gatilho: campo 'Código' ou 'Nome')
                    es_campo_inicio = any(label in texto_celula for label in ['Código', 'Nome do(a) trabalhador(a)'])
                    
                    if es_campo_inicio and dados_atuais:
                        # Se já temos outros campos preenchidos, salva a ficha anterior
                        # (Evita salvar ficha vazia ou duplicar se o gatilho bater várias vezes na mesma página)
                        if len(dados_atuais) > 2: # Mais do que apenas metadados
                            contador_fichas += 1
                            dados_atuais['ficha_n'] = contador_fichas
                            dados_atuais['arquivo_origem'] = os.path.basename(caminho_arquivo)
                            dados_atuais['data_extracao'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                            todas_fichas.append(dados_atuais)
                            dados_atuais = {}

                    # Extração normal (Método 1: Label\nValor)
                    if '\n' in texto_celula:
                        linhas = texto_celula.split('\n')
                        if len(linhas) >= 2:
                            label = linhas[0].strip()
                            valor = '\n'.join(linhas[1:]).strip()
                            
                            if label in campos_mapeamento:
                                campo_chave = campos_mapeamento[label]
                                
                                if campo_chave in campos_endereco:
                                    # Lógica do endereço residencial (index >= 15 da tabela)
                                    if row_idx >= 15:
                                        if campo_chave not in dados_atuais or not dados_atuais[campo_chave]:
                                            dados_atuais[campo_chave] = valor
                                else:
                                    if campo_chave not in dados_atuais or not dados_atuais[campo_chave]:
                                        dados_atuais[campo_chave] = valor
                    
                    # Método 2 (Label embutido)
                    for label, campo_chave in campos_mapeamento.items():
                        if campo_chave in campos_endereco: continue
                        if label in texto_celula and label + '\n' not in texto_celula:
                            partes = texto_celula.split(label, 1)
                            if len(partes) == 2:
                                valor = partes[1].strip().strip('\n').strip()
                                if valor and (campo_chave not in dados_atuais or not dados_atuais[campo_chave]):
                                    dados_atuais[campo_chave] = valor
        
        # Adicionar a última ficha
        if dados_atuais and len(dados_atuais) > 2:
            contador_fichas += 1
            dados_atuais['ficha_n'] = contador_fichas
            dados_atuais['arquivo_origem'] = os.path.basename(caminho_arquivo)
            dados_atuais['data_extracao'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            todas_fichas.append(dados_atuais)
            
        return todas_fichas
    
    def processar_arquivo(self):
        """Processa o arquivo ficha por ficha"""
        try:
            self.adicionar_log("="*60, 'header')
            self.adicionar_log("Iniciando extração inteligente multi-ficha...", 'header')
            self.adicionar_log("="*60, 'header')
            
            arquivo = self.arquivo_selecionado.get()
            self.adicionar_log(f"Arquivo: {os.path.basename(arquivo)}", 'info')
            
            self.status_label.config(text="Analisando documento e separando fichas...")
            self.progress_bar['value'] = 20
            self.percent_label.config(text="20%")
            
            lista_dados = self.extrair_todas_as_fichas(arquivo)
            num_fichas = len(lista_dados)
            
            if num_fichas == 0:
                self.adicionar_log("❌ Nenhuma ficha identificada.", 'error')
                messagebox.showwarning("Aviso", "Nenhum dado foi encontrado no arquivo.")
                return

            self.adicionar_log(f"✓ Identificadas {num_fichas} fichas individuais.", 'success')
            
            # Criar DataFrame e exportar
            self.status_label.config(text="Organizando dados estruturados...")
            self.progress_bar['value'] = 60
            self.percent_label.config(text="60%")
            
            df = pd.DataFrame(lista_dados)
            
            # Reordenar colunas
            colunas_prioritarias = ['ficha_n', 'nome', 'cpf', 'data_admissao', 'contrato', 'funcao']
            cols = [c for c in colunas_prioritarias if c in df.columns]
            cols += [c for c in df.columns if c not in cols]
            df = df[cols]
            
            pasta_saida = os.path.dirname(arquivo)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            arquivo_saida = os.path.join(pasta_saida, f'fichas_organizadas_{timestamp}.xlsx')
            
            # Formatar colunas para evitar a "bagunça"
            df.to_excel(arquivo_saida, index=False, engine='openpyxl')
            
            self.progress_bar['value'] = 100
            self.percent_label.config(text="100%")
            self.status_label.config(text="Concluído!")
            
            self.adicionar_log(f"✓ {num_fichas} funcionários salvos organizadamente.", 'success')
            self.adicionar_log(f"Arquivo gerado: {os.path.basename(arquivo_saida)}", 'success')
            self.adicionar_log("="*60, 'header')
            
            messagebox.showinfo("Sucesso", f"Processamento concluído!\n\nForam extraídas {num_fichas} fichas organizadas.\nArquivo: {os.path.basename(arquivo_saida)}")
            
        except Exception as e:
            self.adicionar_log(f"ERRO: {str(e)}", 'error')
            messagebox.showerror("Erro", f"Ocorreu um problema ao organizar os dados:\n{str(e)}")
        finally:
            self.processando = False
            self.btn_processar.config(state='normal')
            self.progress_bar['value'] = 0
            self.percent_label.config(text="0%")
    
    def sair(self):
        """Fecha a aplicação"""
        if self.processando:
            if messagebox.askyesno("Confirmar", "Há um processamento em andamento. Deseja realmente sair?"):
                self.root.destroy()
        else:
            self.root.destroy()


def main():
    root = tk.Tk()
    app = ExtratorWordArquivoUnico(root)
    root.mainloop()


if __name__ == "__main__":
    main()
